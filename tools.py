from langchain_community.tools import WikipediaQueryRun
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_community.tools import ddg_search
from langchain.tools import Tool
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document


search_tool = Tool(
    func=ddg_search.DuckDuckGoSearchRun().run,
    name="DuckDuckGo_Search",
    description="Search the web using DuckDuckGo"
)

wikipedia_tool = WikipediaQueryRun(
    api_wrapper=WikipediaAPIWrapper(top_k_results=1, doc_content_chars_max=100)
)


def save_to_txt(data: str, filename: str = "research_output.txt"):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    formatted_text = f"--- Research Output ---\nTimestamp: {timestamp}\n\n{data}\n\n"
    with open(filename, "a", encoding="utf-8") as f:
        f.write(formatted_text)
    return f"Data successfully saved to {filename}"

save_tool = Tool(
    func=save_to_txt,
    name="Save_to_Text_File",
    description="Saves the research output to a text file"
)

def save_to_pdf(recipe: dict, filename: str = "recipe.pdf"):
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()
    content = []

    content.append(Paragraph(f"<b>{recipe['topic']}</b>", styles['Title']))
    content.append(Spacer(1, 12))
    content.append(Paragraph("<b>Summary:</b> " + recipe['summary'], styles['Normal']))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Ingredients:</b>", styles['Heading2']))
    for item in recipe['ingredients']:
        content.append(Paragraph(f"- {item}", styles['Normal']))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Instructions:</b>", styles['Heading2']))
    for step in recipe['instructions']:
        content.append(Paragraph(f"{step}", styles['Normal']))
    content.append(Spacer(1, 12))

    if recipe.get('sources'):
        content.append(Paragraph("<b>Sources:</b>", styles['Heading2']))
        for src in recipe['sources']:
            content.append(Paragraph(src, styles['Normal']))

    doc.build(content)
    return f"Recipe saved as PDF: {filename}"

save_pdf_tool = Tool(
    func=save_to_pdf,
    name="Save_to_PDF",
    description="Save the recipe as a PDF file"
)

def save_to_word(recipe: dict, filename: str = "recipe.docx"):
    doc = Document()
    doc.add_heading(recipe['topic'], 0)

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(recipe['summary'])

    doc.add_heading('Ingredients', level=1)
    for item in recipe['ingredients']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Instructions', level=1)
    for step in recipe['instructions']:
        doc.add_paragraph(step, style='List Number')

    if recipe.get('sources'):
        doc.add_heading('Sources', level=1)
        for src in recipe['sources']:
            doc.add_paragraph(src)

    doc.save(filename)
    return f"Recipe saved as Word file: {filename}"

save_word_tool = Tool(
    func=save_to_word,
    name="Save_to_Word",
    description="Save the recipe as a Word document"
)
